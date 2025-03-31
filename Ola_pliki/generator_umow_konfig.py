import pandas as pd
from docx import Document
import re
import os
import argparse
from datetime import datetime

def clean_value(value):
    """Czyści wartość z pliku Excel, usuwając nawiasy kwadratowe i inne znaki specjalne."""
    if pd.isna(value) or value is None:
        return ""
    
    value = str(value).strip()
    # Usuń nawiasy kwadratowe
    value = re.sub(r'\[|\]', '', value)
    return value.strip()

def format_date(date_str):
    """Formatuje datę do formatu DD.MM.YYYY"""
    try:
        if pd.isna(date_str):
            return ""
        # Próbuj różnych formatów daty
        for fmt in ['%Y-%m-%d', '%d.%m.%Y', '%d/%m/%Y', '%Y/%m/%d']:
            try:
                date = datetime.strptime(str(date_str), fmt)
                return date.strftime('%d.%m.%Y')
            except ValueError:
                continue
        return str(date_str)
    except Exception:
        return str(date_str)

def get_gender_suffix(name):
    """Określa płeć na podstawie imienia i zwraca odpowiedni sufiks"""
    female_endings = ['a', 'ia', 'ja', 'ka', 'la', 'ma', 'na', 'ra', 'sa', 'ta', 'wa', 'za']
    if not name:
        return 'y'  # domyślnie męski
    
    # Sprawdź czy imię kończy się na typowe żeńskie końcówki
    for ending in female_endings:
        if name.lower().endswith(ending):
            return 'a'
    return 'y'

def generate_contracts(excel_file, template_file, output_dir, sheet_name, fields, fields_mapping=None):
    """
    Generuje umowy o pracę na podstawie danych z pliku Excel.
    
    Parametry:
    - excel_file: ścieżka do pliku Excel z danymi
    - template_file: ścieżka do szablonu umowy w formacie .docx
    - output_dir: katalog, w którym zostaną zapisane wygenerowane umowy
    - sheet_name: nazwa arkusza w pliku Excel
    - fields: lista nazw wierszy, które zawierają potrzebne dane
    - fields_mapping: słownik do mapowania nazw pól z Excela na nazwy używane w szablonie
    """
    # Utwórz katalog wyjściowy, jeśli nie istnieje
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"Utworzono katalog: {output_dir}")
    
    # Wczytaj arkusz z danymi pracowników
    try:
        df = pd.read_excel(excel_file, sheet_name=sheet_name)
        print(f"Wczytano dane z arkusza '{sheet_name}' z pliku {excel_file}")
        
        # Sprawdź czy dokument Word istnieje
        if not os.path.exists(template_file):
            print(f"Błąd: Nie znaleziono szablonu umowy: {template_file}")
            return 0
        
        # Kolumna z indeksami wierszy (zazwyczaj Unnamed: 0)
        index_col = df.columns[0]
        
        # Przejdź przez każdą kolumnę (pracownika) i wygeneruj umowę
        liczba_umow = 0
        
        for pracownik_col in df.columns[1:]:  # Pomiń pierwszą kolumnę z indeksami
            # Sprawdź, czy kolumna zawiera dane (nie jest całkowicie pusta)
            if df[pracownik_col].notna().sum() > 0:
                print(f"\nGenerowanie umowy dla: {pracownik_col}")
                
                # Przygotuj słownik z danymi pracownika
                dane_pracownika = {}
                
                # Nazwa pracownika (bez nawiasów kwadratowych, jeśli są)
                nazwa_pracownika = clean_value(pracownik_col)
                dane_pracownika["imie_nazwisko"] = nazwa_pracownika
                
                # Pobierz dane z odpowiednich wierszy
                for row_name in fields:
                    try:
                        # Znajdź wiersz, który zawiera dany klucz
                        matching_rows = df[df[index_col].str.contains(row_name, na=False)]
                        if not matching_rows.empty:
                            value = matching_rows.iloc[0][pracownik_col]
                            # Użyj mapowania pól, jeśli istnieje
                            if fields_mapping and row_name in fields_mapping:
                                key = fields_mapping[row_name]
                            else:
                                key = row_name.replace("/", "_").replace(" ", "_")
                            
                            # Formatuj datę jeśli pole zawiera datę
                            if any(date_key in key.lower() for date_key in ['data', 'termin']):
                                value = format_date(value)
                            
                            dane_pracownika[key] = clean_value(value)
                    except Exception as e:
                        print(f"  Błąd podczas pobierania danych z wiersza '{row_name}': {e}")
                
                # Dodaj informację o płci
                imie = nazwa_pracownika.split()[0] if nazwa_pracownika else ""
                gender_suffix = get_gender_suffix(imie)
                dane_pracownika["gender_suffix"] = gender_suffix
                
                # Wyświetl zebrane dane
                print("  Zebrane dane:")
                for key, value in dane_pracownika.items():
                    print(f"  - {key}: {value}")
                
                # Wczytaj szablon dokumentu
                doc = Document(template_file)
                
                # Zastąp dane w dokumencie
                zamienione = 0
                
                # Lista wzorców do wyszukiwania w dokumencie
                patterns = [
                    r'\[.*?\]',         # Tekst w nawiasach kwadratowych
                    r'\{.*?\}',         # Tekst w nawiasach klamrowych
                    r'_+',              # Podkreślniki
                    r'\.{3,}',          # Wielokropek
                    r'<<.*?>>',         # Tekst w podwójnych nawiasach ostrych
                    r'#.*?#',           # Tekst między znakami hash
                ]
                
                # Zamień dane w paragrafach
                for para in doc.paragraphs:
                    text = para.text
                    text_changed = False
                    
                    # Najpierw sprawdź, czy paragraf zawiera jakiekolwiek wzorce do zastąpienia
                    contains_pattern = False
                    for pattern in patterns:
                        if re.search(pattern, text):
                            contains_pattern = True
                            break
                            
                    if contains_pattern:
                        # Zastąp dane w paragrafie
                        for key, value in dane_pracownika.items():
                            if value and len(value) > 0:
                                # Szukaj wartości bezpośrednio
                                if key == "imie_nazwisko":
                                    if "[Aneta Motyl]" in text:
                                        text = text.replace("[Aneta Motyl]", value)
                                        text_changed = True
                                    # Szukaj innych wzorców dla imienia i nazwiska
                                    imie_nazwisko_patterns = [
                                        r'\[.*?imię.*?nazwisko.*?\]',
                                        r'\[.*?Imię.*?Nazwisko.*?\]',
                                        r'\[imię i nazwisko\]',
                                        r'\[Imię i Nazwisko\]'
                                    ]
                                    for pattern in imie_nazwisko_patterns:
                                        if re.search(pattern, text, re.IGNORECASE):
                                            text = re.sub(pattern, value, text, flags=re.IGNORECASE)
                                            text_changed = True
                                else:
                                    # Szukaj wzorca zawierającego klucz
                                    key_patterns = [f"\\[.*?{key}.*?\\]", f"\\{{.*?{key}.*?\\}}"]
                                    for pattern in key_patterns:
                                        if re.search(pattern, text, re.IGNORECASE):
                                            text = re.sub(pattern, value, text, flags=re.IGNORECASE)
                                            text_changed = True
                        
                        # Zastąp wzorce z sufiksem płci
                        gender_patterns = [
                            r'\[zatrudniony/a\]',
                            r'\[pracownik/a\]',
                            r'\[pracowniczy/a\]'
                        ]
                        for pattern in gender_patterns:
                            if re.search(pattern, text, re.IGNORECASE):
                                replacement = f"zatrudnion{gender_suffix}" if "zatrudniony" in pattern.lower() else \
                                           f"pracownik{gender_suffix}" if "pracownik" in pattern.lower() else \
                                           f"pracownicz{gender_suffix}"
                                text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
                                text_changed = True
                    
                    # Jeśli tekst został zmieniony, zaktualizuj paragraf
                    if text_changed:
                        para.text = text
                        zamienione += 1
                
                # Zamień dane w tabelach
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                text = paragraph.text
                                text_changed = False
                                
                                # Najpierw sprawdź, czy paragraf zawiera jakiekolwiek wzorce do zastąpienia
                                contains_pattern = False
                                for pattern in patterns:
                                    if re.search(pattern, text):
                                        contains_pattern = True
                                        break
                                        
                                if contains_pattern:
                                    # Zastąp dane w paragrafie
                                    for key, value in dane_pracownika.items():
                                        if value and len(value) > 0:
                                            # Szukaj wartości bezpośrednio
                                            if key == "imie_nazwisko":
                                                if "[Aneta Motyl]" in text:
                                                    text = text.replace("[Aneta Motyl]", value)
                                                    text_changed = True
                                                # Szukaj innych wzorców dla imienia i nazwiska
                                                imie_nazwisko_patterns = [
                                                    r'\[.*?imię.*?nazwisko.*?\]',
                                                    r'\[.*?Imię.*?Nazwisko.*?\]',
                                                    r'\[imię i nazwisko\]',
                                                    r'\[Imię i Nazwisko\]'
                                                ]
                                                for pattern in imie_nazwisko_patterns:
                                                    if re.search(pattern, text, re.IGNORECASE):
                                                        text = re.sub(pattern, value, text, flags=re.IGNORECASE)
                                                        text_changed = True
                                            else:
                                                # Szukaj wzorca zawierającego klucz
                                                key_patterns = [f"\\[.*?{key}.*?\\]", f"\\{{.*?{key}.*?\\}}"]
                                                for pattern in key_patterns:
                                                    if re.search(pattern, text, re.IGNORECASE):
                                                        text = re.sub(pattern, value, text, flags=re.IGNORECASE)
                                                        text_changed = True
                                
                                    # Szukaj konkretnych wartości z nawiasami
                                    if "[20.01.2000]" in text and "data_urodzenia" in dane_pracownika:
                                        text = text.replace("[20.01.2000]", dane_pracownika["data_urodzenia"])
                                        text_changed = True
                                    if "[05.12.2024]" in text and "data_zawarcia_umowy" in dane_pracownika:
                                        text = text.replace("[05.12.2024]", dane_pracownika["data_zawarcia_umowy"])
                                        text_changed = True
                                    if "[01.01.2025]" in text and "początek_umowy_dzień_rozpoczęcia_pracy" in dane_pracownika:
                                        text = text.replace("[01.01.2025]", dane_pracownika["początek_umowy_dzień_rozpoczęcia_pracy"])
                                        text_changed = True
                                    if "[31.03.2025]" in text and "koniec_umowy" in dane_pracownika:
                                        text = text.replace("[31.03.2025]", dane_pracownika["koniec_umowy"])
                                        text_changed = True
                                    if "[Asystent]" in text and "stanowisko" in dane_pracownika:
                                        text = text.replace("[Asystent]", dane_pracownika["stanowisko"])
                                        text_changed = True
                                
                                # Jeśli tekst został zmieniony, zaktualizuj paragraf
                                if text_changed:
                                    paragraph.text = text
                                    zamienione += 1
                
                # Utwórz nazwe pliku wyjściowego
                nazwa_pliku = f"{nazwa_pracownika.replace(' ', '_')}_umowa_okres_probny.docx"
                plik_wyjsciowy = os.path.join(output_dir, nazwa_pliku)
                
                # Zapisz dokument
                doc.save(plik_wyjsciowy)
                print(f"  Zastąpiono {zamienione} miejsc w dokumencie.")
                print(f"  Zapisano umowę: {plik_wyjsciowy}")
                liczba_umow += 1
        
        return liczba_umow
            
    except Exception as e:
        print(f"Wystąpił błąd: {e}")
        return 0

def main():
    parser = argparse.ArgumentParser(description="Generator umów o pracę na podstawie danych z Excela")
    
    parser.add_argument("--excel", "-e", type=str, default="Nadgodziny_nieobecności_pracownicy.xlsx",
                        help="Ścieżka do pliku Excel z danymi pracowników")
    parser.add_argument("--template", "-t", type=str, default="UMOWA O PRACĘ NA OKRES PRÓBNY.docx",
                        help="Ścieżka do szablonu umowy w formacie .docx")
    parser.add_argument("--output", "-o", type=str, default="Wygenerowane_umowy",
                        help="Katalog, w którym zostaną zapisane wygenerowane umowy")
    parser.add_argument("--sheet", "-s", type=str, default="pracownicy",
                        help="Nazwa arkusza w pliku Excel")
    parser.add_argument("--fields", type=str, nargs="+", 
                        default=["data urodzenia", "data zawarcia umowy", "początek umowy/dzień rozpoczęcia pracy", 
                                 "koniec umowy", "stanowisko"],
                        help="Nazwy wierszy, które zawierają potrzebne dane")
    
    args = parser.parse_args()
    
    print("Generator umów o pracę na podstawie danych z Excela")
    print("--------------------------------------------------")
    print(f"Plik Excel: {args.excel}")
    print(f"Szablon umowy: {args.template}")
    print(f"Katalog wyjściowy: {args.output}")
    print(f"Arkusz: {args.sheet}")
    print(f"Wiersze z danymi: {', '.join(args.fields)}")
    print()
    
    # Definicja mapowania nazw pól (można dostosować według potrzeb)
    fields_mapping = {
        "data urodzenia": "data_urodzenia",
        "data zawarcia umowy": "data_zawarcia_umowy",
        "początek umowy/dzień rozpoczęcia pracy": "data_rozpoczecia_pracy",
        "koniec umowy": "data_zakonczenia_umowy",
        "stanowisko": "stanowisko"
    }
    
    liczba_umow = generate_contracts(
        excel_file=args.excel,
        template_file=args.template,
        output_dir=args.output,
        sheet_name=args.sheet,
        fields=args.fields,
        fields_mapping=fields_mapping
    )
    
    print(f"\nWygenerowano {liczba_umow} umów w katalogu {args.output}")

if __name__ == "__main__":
    main() 