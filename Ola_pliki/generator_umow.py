import pandas as pd
from docx import Document
import re
import os
from datetime import datetime

def clean_value(value):
    """Czyści wartość z pliku Excel, usuwając nawiasy kwadratowe i inne znaki specjalne."""
    if pd.isna(value) or value is None:
        return ""
    
    value = str(value).strip()
    # Usuń nawiasy kwadratowe
    value = re.sub(r'\[|\]', '', value)
    return value.strip()

def main():
    print("Generator umów o pracę na podstawie danych z Excela")
    print("--------------------------------------------------")
    
    # Konfiguracja ścieżek plików
    excel_file = "Nadgodziny_nieobecności_pracownicy.xlsx"
    template_file = "UMOWA O PRACĘ NA OKRES PRÓBNY.docx"
    output_dir = "Wygenerowane_umowy"
    
    # Utwórz katalog wyjściowy, jeśli nie istnieje
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"Utworzono katalog: {output_dir}")
    
    # Wczytaj arkusz z danymi pracowników
    try:
        df_pracownicy = pd.read_excel(excel_file, sheet_name="pracownicy")
        print(f"Wczytano dane pracowników z pliku {excel_file}")
        
        # Sprawdź czy dokument Word istnieje
        if not os.path.exists(template_file):
            print(f"Błąd: Nie znaleziono szablonu umowy: {template_file}")
            return
        
        # Kolumna z indeksami wierszy (zazwyczaj Unnamed: 0)
        index_col = df_pracownicy.columns[0]
        
        # Wiersze, które nas interesują
        interesting_rows = [
            "data urodzenia",
            "data zawarcia umowy",
            "początek umowy/dzień rozpoczęcia pracy",
            "koniec umowy",
            "stanowisko",
            # Możesz dodać więcej wierszy, które są potrzebne
        ]
        
        # Przejdź przez każdą kolumnę (pracownika) i wygeneruj umowę
        liczba_umow = 0
        
        for pracownik_col in df_pracownicy.columns[1:]:  # Pomiń pierwszą kolumnę z indeksami
            # Sprawdź, czy kolumna zawiera dane (nie jest całkowicie pusta)
            if df_pracownicy[pracownik_col].notna().sum() > 0:
                print(f"\nGenerowanie umowy dla: {pracownik_col}")
                
                # Przygotuj słownik z danymi pracownika
                dane_pracownika = {}
                
                # Nazwa pracownika (bez nawiasów kwadratowych, jeśli są)
                nazwa_pracownika = clean_value(pracownik_col)
                dane_pracownika["imie_nazwisko"] = nazwa_pracownika
                
                # Pobierz dane z odpowiednich wierszy
                for row_name in interesting_rows:
                    try:
                        # Znajdź wiersz, który zawiera dany klucz
                        matching_rows = df_pracownicy[df_pracownicy[index_col].str.contains(row_name, na=False)]
                        if not matching_rows.empty:
                            value = matching_rows.iloc[0][pracownik_col]
                            key = row_name.replace("/", "_").replace(" ", "_")
                            dane_pracownika[key] = clean_value(value)
                    except Exception as e:
                        print(f"  Błąd podczas pobierania danych z wiersza '{row_name}': {e}")
                
                # Wyświetl zebrane dane
                print("  Zebrane dane:")
                for key, value in dane_pracownika.items():
                    print(f"  - {key}: {value}")
                
                # Wczytaj szablon dokumentu
                doc = Document(template_file)
                
                # Zastąp dane w dokumencie
                zamienione = 0
                
                # Zamień dane w paragrafach
                for para in doc.paragraphs:
                    for key, value in dane_pracownika.items():
                        if value and len(value) > 0:
                            # Szukaj wartości w nawiasach kwadratowych lub klamrowych
                            patterns = [
                                r'\[.*?' + key + r'.*?\]',          # [*key*]
                                r'\{.*?' + key + r'.*?\}',          # {*key*}
                                r'\[.*?\]',                         # Dowolne nawiasy kwadratowe
                                r'\{.*?\}',                         # Dowolne nawiasy klamrowe
                            ]
                            
                            # Jeśli to imię i nazwisko, szukaj konkretnie tego wzorca
                            if key == "imie_nazwisko":
                                patterns.append(r'\[.*?imię.*?nazwisko.*?\]')
                                patterns.append(r'\[.*?Imię.*?Nazwisko.*?\]')
                            
                            for pattern in patterns:
                                if re.search(pattern, para.text, re.IGNORECASE):
                                    para.text = re.sub(pattern, value, para.text, flags=re.IGNORECASE)
                                    zamienione += 1
                                    
                            # Również wyszukaj konkretne wartości do zamiany (np. [Aneta Motyl])
                            if key == "imie_nazwisko":
                                para.text = para.text.replace("[Aneta Motyl]", value)
                            elif key == "data_urodzenia":
                                para.text = para.text.replace("[20.01.2000]", value)
                            elif key == "data_zawarcia_umowy":
                                para.text = para.text.replace("[05.12.2024]", value)
                            elif key == "stanowisko":
                                para.text = para.text.replace("[Asystent]", value)
                            elif key == "początek_umowy_dzień_rozpoczęcia_pracy":
                                para.text = para.text.replace("[01.01.2025]", value)
                            elif key == "koniec_umowy":
                                para.text = para.text.replace("[31.03.2025]", value)
                
                # Zamień dane w tabelach
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for key, value in dane_pracownika.items():
                                    if value and len(value) > 0:
                                        # Szukaj wartości w nawiasach kwadratowych lub klamrowych
                                        patterns = [
                                            r'\[.*?' + key + r'.*?\]',       # [*key*]
                                            r'\{.*?' + key + r'.*?\}',       # {*key*}
                                            r'\[.*?\]',                      # Dowolne nawiasy kwadratowe
                                            r'\{.*?\}',                      # Dowolne nawiasy klamrowe
                                        ]
                                        
                                        # Jeśli to imię i nazwisko, szukaj konkretnie tego wzorca
                                        if key == "imie_nazwisko":
                                            patterns.append(r'\[.*?imię.*?nazwisko.*?\]')
                                            patterns.append(r'\[.*?Imię.*?Nazwisko.*?\]')
                                        
                                        for pattern in patterns:
                                            if re.search(pattern, paragraph.text, re.IGNORECASE):
                                                paragraph.text = re.sub(pattern, value, paragraph.text, flags=re.IGNORECASE)
                                                zamienione += 1
                                                
                                        # Również wyszukaj konkretne wartości do zamiany (np. [Aneta Motyl])
                                        if key == "imie_nazwisko":
                                            paragraph.text = paragraph.text.replace("[Aneta Motyl]", value)
                                        elif key == "data_urodzenia":
                                            paragraph.text = paragraph.text.replace("[20.01.2000]", value)
                                        elif key == "data_zawarcia_umowy":
                                            paragraph.text = paragraph.text.replace("[05.12.2024]", value)
                                        elif key == "stanowisko":
                                            paragraph.text = paragraph.text.replace("[Asystent]", value)
                                        elif key == "początek_umowy_dzień_rozpoczęcia_pracy":
                                            paragraph.text = paragraph.text.replace("[01.01.2025]", value)
                                        elif key == "koniec_umowy":
                                            paragraph.text = paragraph.text.replace("[31.03.2025]", value)
                
                # Utwórz nazwe pliku wyjściowego
                nazwa_pliku = f"{nazwa_pracownika.replace(' ', '_')}_umowa_okres_probny.docx"
                plik_wyjsciowy = os.path.join(output_dir, nazwa_pliku)
                
                # Zapisz dokument
                doc.save(plik_wyjsciowy)
                print(f"  Zastąpiono {zamienione} miejsc w dokumencie.")
                print(f"  Zapisano umowę: {plik_wyjsciowy}")
                liczba_umow += 1
        
        print(f"\nWygenerowano {liczba_umow} umów w katalogu {output_dir}")
            
    except Exception as e:
        print(f"Wystąpił błąd: {e}")

if __name__ == "__main__":
    main() 