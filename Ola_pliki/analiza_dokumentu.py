from docx import Document
import os

# Analiza pliku .docx
docx_file = "UMOWA O PRACĘ NA OKRES PRÓBNY.docx"
print(f"--- SZCZEGÓŁOWA ANALIZA DOKUMENTU: {docx_file} ---")

try:
    doc = Document(docx_file)
    
    print(f"Dokument zawiera {len(doc.paragraphs)} paragrafów i {len(doc.tables)} tabel.\n")
    
    # Analizuj zawartość wszystkich paragrafów
    print("ZAWARTOŚĆ PARAGRAFÓW:")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():  # Pomijamy puste paragrafy
            print(f"Paragraf {i}: {para.text}")
            
            # Sprawdź styl paragrafu
            if para.style:
                print(f"   Styl: {para.style.name}")
            
            # Sprawdź formatowanie tekstu w paragrafie
            if para.runs:
                formats = []
                for run in para.runs:
                    if run.bold: formats.append("pogrubiony")
                    if run.italic: formats.append("kursywa")
                    if run.underline: formats.append("podkreślony")
                if formats:
                    print(f"   Formatowanie: {', '.join(formats)}")
            print()
    
    # Analizuj zawartość tabel
    if doc.tables:
        print("\nZAWARTOŚĆ TABEL:")
        for i, table in enumerate(doc.tables):
            print(f"Tabela {i+1} ({len(table.rows)} wierszy x {len(table.columns)} kolumn):")
            
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        print(f"   Komórka [{row_idx},{col_idx}]: {cell_text}")
            print()
    
    # Sugestie miejsc do uzupełnienia
    print("\nSUGEROWANE MIEJSCA DO UZUPEŁNIENIA:")
    print("Na podstawie analizy dokumentu, poniżej znajdują się potencjalne miejsca do uzupełnienia:")
    
    # Lista typowych pól do uzupełnienia w umowie o pracę
    typical_fields = [
        "data",
        "imię i nazwisko",
        "adres",
        "stanowisko",
        "wymiar czasu pracy",
        "wynagrodzenie",
        "data rozpoczęcia pracy",
        "miejsce wykonywania pracy",
        "okres próbny"
    ]
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.lower()
        for field in typical_fields:
            if field in text:
                print(f"- Paragraf {i}: zawiera informację o '{field}' - możliwe miejsce do wstawienia danych.")
                
    # Sprawdź tabele pod kątem miejsc do uzupełnienia
    for i, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                text = cell.text.lower()
                for field in typical_fields:
                    if field in text:
                        print(f"- Tabela {i+1}, komórka [{row_idx},{col_idx}]: zawiera informację o '{field}' - możliwe miejsce do wstawienia danych.")
except Exception as e:
    print(f"Błąd podczas analizy dokumentu Word: {e}")

print("\n--- PROPOZYCJA ROZWIĄZANIA ---")
print("Na podstawie analizy dokumentu Word i pliku Excel, proponuję stworzenie skryptu, który:")
print("1. Odczyta dane pracowników z arkusza 'pracownicy' w pliku Excel")
print("2. Dla każdego pracownika (kolumny) znajdzie dane w odpowiednich wierszach")
print("3. Wstawi te dane w odpowiednie miejsca w dokumencie Word")
print("4. Zapisze wygenerowany dokument jako '[Imię]_[Nazwisko]_umowa.docx'") 