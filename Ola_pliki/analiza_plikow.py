import pandas as pd
import os
from docx import Document
import re

# Wyświetl zawartość katalogu
print("Dostępne pliki w katalogu:")
for file in os.listdir():
    print(f"- {file}")

# Analiza pliku Excel
excel_file = "Nadgodziny_nieobecności_pracownicy.xlsx"
print(f"\n--- ANALIZA PLIKU EXCEL: {excel_file} ---")

try:
    # Sprawdź, jakie arkusze są dostępne w pliku Excel
    xlsx = pd.ExcelFile(excel_file)
    print(f"Dostępne arkusze: {xlsx.sheet_names}")
    
    # Przeanalizuj każdy arkusz
    for sheet_name in xlsx.sheet_names:
        print(f"\n-- Arkusz: {sheet_name} --")
        df = pd.read_excel(excel_file, sheet_name=sheet_name)
        
        # Wyświetl pierwsze kilka wierszy
        print(f"Pierwsze wiersze danych:")
        print(df.head())
        
        # Wyświetl nazwy kolumn
        print(f"Kolumny w arkuszu:")
        for col in df.columns:
            print(f"- {col}")
except Exception as e:
    print(f"Błąd podczas odczytu pliku Excel: {e}")

# Analiza pliku .docx
docx_file = "UMOWA O PRACĘ NA OKRES PRÓBNY.docx"
print(f"\n--- ANALIZA PLIKU WORD: {docx_file} ---")

try:
    doc = Document(docx_file)
    
    print(f"Liczba paragrafów: {len(doc.paragraphs)}")
    
    # Szukaj potencjalnych miejsc do uzupełnienia (np. tekst w nawiasach, podkreślniki, etc.)
    placeholders = []
    
    # Wzorce do wyszukiwania możliwych miejsc do uzupełnienia
    patterns = [
        r'\[.*?\]',         # Tekst w nawiasach kwadratowych
        r'\{.*?\}',         # Tekst w nawiasach klamrowych
        r'_+',              # Podkreślniki
        r'\.{3,}',          # Wielokropek
        r'<<.*?>>',         # Tekst w podwójnych nawiasach ostrych
        r'#.*?#',           # Tekst między znakami hash
    ]
    
    # Przeszukaj wszystkie paragrafy
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue  # Pomijaj puste paragrafy
            
        # Szukaj wzorców
        for pattern in patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                placeholders.append({
                    'tekst': match,
                    'paragraf_nr': i,
                    'paragraf': text[:100] + ('...' if len(text) > 100 else '')
                })
    
    # Wyświetl znalezione miejsca do uzupełnienia
    if placeholders:
        print(f"\nZnaleziono {len(placeholders)} potencjalnych miejsc do uzupełnienia:")
        for i, p in enumerate(placeholders, 1):
            print(f"{i}. '{p['tekst']}' w paragrafie {p['paragraf_nr']} - '{p['paragraf']}'")
    else:
        print("\nNie znaleziono wyraźnie oznaczonych miejsc do uzupełnienia.")
        
    # Wyświetl pierwsze kilka paragrafów jako przykład
    print("\nPierwsze paragrafy dokumentu:")
    for i, para in enumerate(doc.paragraphs[:10]):
        if para.text.strip():
            print(f"Paragraf {i}: {para.text[:100]}")
            
except Exception as e:
    print(f"Błąd podczas odczytu pliku Word: {e}")

# Podsumowanie
print("\n--- PODSUMOWANIE ANALIZY ---")
print("Na podstawie analizy plików możemy stworzyć automatyzację, która:")
print("1. Odczyta dane pracowników z arkusza Excel")
print("2. Dla każdego pracownika wygeneruje umowę o pracę, uzupełniając odpowiednie pola")
print("3. Zapisze wygenerowane dokumenty jako indywidualne pliki .docx") 